import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
        left={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
        right={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if not create
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def convert_html_to_docx(html_path, docx_path):
    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    soup = BeautifulSoup(html_content, 'html.parser')
    doc = Document()

    # Set default style to Times New Roman
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    
    # Page setup
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title Page Logic
    title_div = soup.find('div', class_='title-page')
    if title_div:
        for child in title_div.find_all('div'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(child.get_text().strip())
            if 'project-title' in child.get('class', []):
                run.bold = True
                run.font.size = Pt(28)
            else:
                run.font.size = Pt(18)
            # Add spacing
            p.paragraph_format.space_after = Pt(20)
        doc.add_page_break()

    # Content Logic
    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'div', 'table', 'br']):
        if element.name == 'h1':
            # Handle page breaks before h1 (except possibly the first one if already broken)
            h = doc.add_heading(element.get_text().strip(), level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Reset font for headings
            run = h.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(24)
            run.font.color.rgb = None
            run.bold = True
            h.paragraph_format.space_before = Pt(30)
            h.paragraph_format.space_after = Pt(20)
            
        elif element.name == 'h2':
            h = doc.add_heading(element.get_text().strip(), level=2)
            run = h.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(20)
            run.font.color.rgb = None
            run.bold = True
            h.paragraph_format.space_before = Pt(20)
            
        elif element.name == 'h3':
            h = doc.add_heading(element.get_text().strip(), level=3)
            run = h.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.font.color.rgb = None
            run.bold = True
            
        elif element.name == 'p':
            # Skip empty paragraphs or those already handled in title page
            if not element.get_text().strip() or element.find_parent('div', class_='title-page'):
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.first_line_indent = Inches(0.5)
            
            # Check for bold tag inside p
            if element.find('b'):
                for child in element.children:
                    if child.name == 'b':
                        run = p.add_run(child.get_text())
                        run.bold = True
                    else:
                        p.add_run(str(child))
            else:
                p.add_run(element.get_text().strip())
                
        elif element.name == 'div' and 'list-item' in element.get('class', []):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.line_spacing = 2.0
            p.add_run(element.get_text().strip().replace('• ', ''))
            
        elif element.name == 'div' and 'sig-line' in element.get('class', []):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(50)
            run = p.add_run(element.get_text().strip())
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        elif element.name == 'table':
            rows = element.find_all('tr')
            if not rows: continue
            
            table = doc.add_table(rows=0, cols=len(rows[0].find_all(['th', 'td'])))
            table.style = 'Table Grid'
            
            for row in rows:
                cells = row.find_all(['th', 'td'])
                row_cells = table.add_row().cells
                for i, cell in enumerate(cells):
                    row_cells[i].text = cell.get_text().strip()
                    # Bold header cells
                    if cell.name == 'th':
                        for paragraph in row_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                
        elif element.name == 'br' and 'page-break-before:always' in str(element.get('style', '')):
            doc.add_page_break()

    doc.save(docx_path)
    print(f"Successfully created {docx_path}")

if __name__ == "__main__":
    html_file = r"c:\Users\makin\Desktop\project_400l\gadget-installments\details\documentation.doc"
    docx_file = r"c:\Users\makin\Desktop\project_400l\gadget-installments\details\documentation.docx"
    convert_html_to_docx(html_file, docx_file)
